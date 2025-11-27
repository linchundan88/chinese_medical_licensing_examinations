'''
A1, A2
A3 142～144题共用题干
B1 型选择题  (145～146题共用备选答案)

2024 2023 2022 2021 2020
第一单元  A1 型选择题(1～66题)，A2型选择题(67～138题)， A3/A4 型选择题(139～150题)
第二单元   A1 型选择题(1～34题)， A2 型选择题(35～113题)， A3/A4 型选择题(114～144题) ，B1 型选择题(145～150题)
第三单元  A1 型选择题(1～27题)， A2 型选择题(28～114题)， A3/A4 型选择题(115～144题)，B1 型选择题(145～150题)
第四单元  A1 型选择题(1～31题)， A2 型选择题(32～108题)， A3/A4 型选择题(109～142题)  B1 型选择题(143～150题)
参考答案及解析
107.C
124～125.AD
'''
import docx
import re
import pickle
from pathlib import Path


if __name__ == '__main__':
    path_docx = Path(__file__).resolve().parent.parent / 'datafiles' / 'clinical_practice_doctor_examination.docx'
    doc = docx.Document(str(path_docx))

    file_pkl = Path(__file__).resolve().parent.parent / 'datafiles' / 'question_answer_data.pkl'

    list_question_answer = []
    list_correct_answer = []

    list_years = [2024, 2023, 2022, 2021, 2020]
    list_units = ['第一单元', '第二单元', '第三单元', '第四单元']

    mode_question_answer = 'question'
    current_unit = 1

    for index, para in enumerate(doc.paragraphs):
        para_text = para.text.replace(' ', '')
        if para_text == '':
            continue

        if '年临床执业医师资格考试机考真题精编' in para_text:
            match = re.search(r'\d{4}年临床执业医师资格考试机考真题精编', para.text)
            if match:
                year = match.group().replace('年临床执业医师资格考试机考真题精编', '')
                current_year = int(year)
                print(current_year)
                mode_question_answer = 'question'
            else:
                raise Exception(f'error:{para_text}')

        if para_text == '参考答案及解析':
            mode_question_answer = 'answer'
            continue

        is_unit = False
        for index_unit, tmp_unit in enumerate(list_units):
            if tmp_unit in para_text:
                current_unit = index_unit + 1
                is_unit = True
                break
        if is_unit:
            continue

        print(para_text)

        if '关于强直性脊柱炎和类风湿性关节炎的特点' in para_text:
            print('aaa')

        if mode_question_answer == 'question':
            if 'A1型选择题' in para_text:
                question_type = 'A1'
            if 'A2型选择题' in para_text:
                question_type = 'A2'
            if 'A3/A4型选择题' in para_text:
                question_type = 'A3/A4'
            if 'B1型选择题' in para_text:
                question_type = 'B1'

            if 'question_type' not in locals():
                continue

            if question_type in ['A1', 'A2']:
                match = re.search(r'^\d{1,3}\.', para_text)
                if match:
                    match = re.search(r'^\d{1,3}', para_text)
                    tmp_no = int(match.group())
                    if 'current_question_no' in locals():
                        if tmp_no != current_question_no+1 and tmp_no != 1:  # 0.5 cm
                            continue

                    if para_text.startswith('1.5g/'):  # 1.5g/   日，血肌酐90μmol/L 。B 超示双肾大小正常。
                        continue

                    current_question_no = tmp_no
                    current_question_stem = para_text[len(match.group())+1:]

                    current_question_options = ''
                    found_option = False
                    for i in range(1, 10):
                        test1 = doc.paragraphs[index + i].text.replace(' ', '')
                        if test1 == '':
                            continue

                        if test1.startswith('A.'):
                            found_option = True

                        if found_option:
                            current_question_options = current_question_options + test1
                            if test1.startswith('E.'):
                                break
                        else:
                            current_question_stem += test1

                    current_question_options = re.sub(r'B\.', ' B.', current_question_options, 1)
                    current_question_options = re.sub(r'C\.', ' C.', current_question_options, 1)
                    current_question_options = re.sub(r'D\.', ' D.', current_question_options, 1)
                    current_question_options = re.sub(r'E\.', ' E.', current_question_options, 1)
                    print(current_question_stem, current_question_options)

                    dict1 = {'year': current_year, 'unit': current_unit, 'question_type': question_type, 'question_no': current_question_no,
                             'question_stem': current_question_stem, 'question_options': current_question_options}
                    list_question_answer.append(dict1)
                    continue

            if question_type == 'A3/A4':
                match = re.search(r'\d{1,3}～\d{1,3}题共用题干', para_text)  #(108～110题共用题干)  #r'^\(\d{1,3}～\d{1,3}题共用题干'
                if match:
                    questsion_no1, questsion_no2 = match.group().replace('题共用题干', '').split('～')
                    current_question_stem = doc.paragraphs[index+1].text
                    continue

                match = re.search(r'^\d{1,3}\.', para_text)
                if match:
                    current_question_no = int(match.group()[:-1])
                    current_question_stem1 = current_question_stem + para_text[len(match.group())+1:]

                    current_question_options = ''
                    for i in range(1, 10):
                        test1 = doc.paragraphs[index + i].text.replace(' ', '')
                        if test1 == '':
                            continue

                        current_question_options = current_question_options + test1
                        if test1.startswith('E.'):
                            break

                    current_question_options = re.sub(r'B\.', ' B.', current_question_options, 1)
                    current_question_options = re.sub(r'C\.', ' C.', current_question_options, 1)
                    current_question_options = re.sub(r'D\.', ' D.', current_question_options, 1)
                    current_question_options = re.sub(r'E\.', ' E.', current_question_options, 1)

                    dict1 = {'year': current_year, 'unit': current_unit, 'question_type': question_type, 'question_no': current_question_no,
                             'question_stem': current_question_stem1, 'question_options': current_question_options}
                    list_question_answer.append(dict1)
                    continue

            if question_type == 'B1':
                match = re.search(r'^\d{1,3}\.', para_text)
                if match:
                    match = re.search(r'^\d{1,3}', para_text)
                    current_question_no = int(match.group())
                    current_question_stem = para_text[len(match.group()) + 1:]

                    dict1 = {'year': current_year, 'unit': current_unit, 'question_type': question_type, 'question_no': current_question_no,
                             'question_stem': current_question_stem, 'question_options': current_question_options}
                    list_question_answer.append(dict1)

                match = re.search(r'\d{1,3}～\d{1,3}题共用备选答案', para_text)
                if match:
                    questsion_no1, questsion_no2 = match.group().replace('题共用备选答案', '').split('～')
                    current_question_options = ''
                    for i in range(1, 10):
                        test1 = doc.paragraphs[index + i].text.replace(' ', '')
                        if test1 == '':
                            continue

                        current_question_options = current_question_options + test1
                        if test1.startswith('E.'):
                            break

                    current_question_options = re.sub(r'B\.', ' B.', current_question_options, 1)
                    current_question_options = re.sub(r'C\.', ' C.', current_question_options, 1)
                    current_question_options = re.sub(r'D\.', ' D.', current_question_options, 1)
                    current_question_options = re.sub(r'E\.', ' E.', current_question_options, 1)
                    continue

        if mode_question_answer == 'answer':
            # 125.A
            print('year:', current_year, 'unit:', current_unit, para_text)
            match = re.search(r'^\d{1,3}\.[A-E]', para_text)
            if match:
                match = re.search(r'^\d{1,3}', para_text)
                question_no = int(match.group())
                correct_answer = para_text[len(match.group())+1:len(match.group())+2]

                dict1 = {'year': current_year, 'unit': current_unit, 'question_no': question_no, 'correct_answer': correct_answer}
                list_correct_answer.append(dict1)
                continue

            # The explanation of answer is too long to fit in one line.
            match = re.search(r'^\d{1,3}', para_text)
            if not match:
                continue

            # 117～118.AA  130～131.CD  145～147.BBC  2024 unit 3 135~136.AA
            match = re.search(r'^\d{1,3}[～,~]\d{1,3}\.[A-E]+', para_text)
            if match:
                text1 = match.group()
                match = re.search(r'^\d{1,3}[～,~]\d{1,3}', para_text)
                if '～' in match.group():
                    character1 = '～'
                if '~' in match.group():
                    character1 = '~'
                question_no_start, question_no_end = int(match.group().split(character1)[0]), int(match.group().split(character1)[1])
                answers = text1[len(match.group())+1:]
                for tmp_question_no in range(question_no_start, question_no_end+1):
                    correct_answer = answers[tmp_question_no-question_no_start]
                    dict1 = {'year': current_year, 'unit': current_unit, 'question_no': tmp_question_no, 'correct_answer': correct_answer}
                    list_correct_answer.append(dict1)
                continue


    with open(file_pkl, 'wb') as f:
        pickle.dump({'questions': list_question_answer, 'answers': list_correct_answer}, f)
    print('Data saved to question_answer_data.pkl')


    print('OK.')